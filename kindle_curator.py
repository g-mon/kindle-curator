import re
from dataclasses import dataclass
from typing import List, Optional, Tuple
from io import BytesIO

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


TRUNC_PHRASE = "Some highlights have been hidden or truncated due to export limits."


@dataclass
class Entry:
    marker_kind: Optional[str]   # "Page" or "Location" or None
    marker_value: Optional[int]  # numeric for sorting/thresholds
    highlight: str
    note: Optional[str] = None
    truncated: bool = False


@dataclass
class ChapterMark:
    marker_kind: str            # "Page" or "Location"
    marker_value: int
    title: str                  # Chapter title to insert


# ---- Parsing ----
META_LINE_RE = re.compile(
    r"""(?ix)^\s*(
        options |
        added\s+on\s+.* |
        =+\s*$
    )\s*$"""

)

HIGHLIGHT_HEADER_RE = re.compile(
    r"""(?ix)^\s*
    (yellow|blue|pink|orange)\s+highlight
    \s*\|\s*
    (page|location)\s*:\s*(\d+)
    \s*$"""
)

# Dates like: "January, 1st 1925" or "January 1st 1925"
DATE_STAMP_RE = re.compile(
    r"""(?ix)^\s*
    (january|february|march|april|may|june|july|august|september|october|november|december)
    \s*,?\s*
    \d{1,2}(?:st|nd|rd|th)?
    \s+\d{4}
    \s*$"""
)

NOTE_LINE_RE = re.compile(r"(?i)^\s*note\s*:\s*(.*)$")
ELLIPSIS_END_RE = re.compile(r"(…|\.\.\.)\s*$")
TRUNC_PHRASE = "Some highlights have been hidden or truncated due to export limits."


PAGE_RE = re.compile(r"(?i)\bpage\s*[:#]?\s*(\d+)\b")
LOC_RE  = re.compile(r"(?i)\blocation\s*[:#]?\s*(\d+)\b")
NOTE_PREFIX_RE = re.compile(r"(?i)^\s*note\s*[:\-]\s*")

def _contains_trunc_phrase(text: str) -> bool:
    return TRUNC_PHRASE.lower() in (text or "").lower()

def _clean_lines(raw: str) -> List[str]:
    out = []
    for line in raw.splitlines():
        l = line.strip()
        if not l:
            out.append("")
            continue
        # IMPORTANT: do NOT strip TRUNC_PHRASE here; it may appear inside highlights
        if META_LINE_RE.match(l):
            continue
        out.append(l)
    return out


def _split_blocks(lines: List[str]) -> List[List[str]]:
    blocks, buf = [], []
    for l in lines:
        if l == "":
            if buf:
                blocks.append(buf)
                buf = []
        else:
            buf.append(l)
    if buf:
        blocks.append(buf)
    return blocks


def _marker_from_block(block: List[str], last_marker: Tuple[Optional[str], Optional[int]]):
    kind, val = last_marker
    for l in block:
        m = PAGE_RE.search(l)
        if m:
            return ("Page", int(m.group(1)))
        m = LOC_RE.search(l)
        if m:
            return ("Location", int(m.group(1)))
    return (kind, val)


def _strip_marker_fragments(text: str) -> str:
    text = PAGE_RE.sub("", text)
    text = LOC_RE.sub("", text)
    return text.strip(" -|")


def parse_kindle(raw: str) -> List[Entry]:
    lines = _clean_lines(raw)

    entries: List[Entry] = []
    current: Optional[Entry] = None

    def flush():
        nonlocal current
        if not current:
            return

        # Strip truncation phrase from highlight if present
        if TRUNC_PHRASE.lower() in current.highlight.lower():
            current.truncated = True
            current.highlight = re.sub(
                re.escape(TRUNC_PHRASE), "", current.highlight, flags=re.IGNORECASE
            ).strip()

        # If highlight ends with ellipsis, likely truncated
        if ELLIPSIS_END_RE.search(current.highlight.strip() or ""):
            current.truncated = True

        # Keep entry if it has highlight text OR is flagged truncated (even if empty)
        if current.highlight.strip() or current.truncated:
            entries.append(current)

        current = None

    for line in lines:
        l = line.strip()

        if not l:
            continue

        # Start of a new highlight entry
        m = HIGHLIGHT_HEADER_RE.match(l)
        if m:
            flush()
            kind = "Page" if m.group(2).lower() == "page" else "Location"
            val = int(m.group(3))
            current = Entry(
                marker_kind=kind,
                marker_value=val,
                highlight="",
                note=None,
                truncated=False
            )
            continue

        # Ignore junk metadata lines
        if META_LINE_RE.match(l):
            continue

        # Remove standalone date stamps like "January, 1st 1925"
        if DATE_STAMP_RE.match(l):
            continue

        # If we see the truncation phrase as its own line, flag the current entry and skip the line
        if current is not None and TRUNC_PHRASE.lower() in l.lower():
            current.truncated = True
            continue

        # Notes attach to current entry
        nm = NOTE_LINE_RE.match(l)
        if nm and current is not None:
            note_text = nm.group(1).strip()
            if note_text:
                if current.note:
                    current.note += "\n" + note_text
                else:
                    current.note = note_text
            continue

        # Otherwise it's highlight text
        if current is not None:
            if current.highlight:
                current.highlight += "\n" + l
            else:
                current.highlight = l
        else:
            # ignore anything before the first highlight header
            continue

    flush()
    return entries




# ---- DOCX generation ----
def _set_para_base(p):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _add_run(p, text: str, font_name: str, size_pt: int, bold: bool = False, italic: bool = False):
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = Pt(size_pt)
    r.bold = bold
    r.italic = italic
    return r


def build_docx(
    title: str,
    entries: List[Entry],
    reading_note: Optional[str],
    chapters: List[ChapterMark],
    font_name: str = "Calibri"
) -> bytes:
    doc = Document()

    # Normal style: body 10pt
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10)

    # Title 12pt bold
    p_title = doc.add_paragraph()
    _set_para_base(p_title)
    _add_run(p_title, title, font_name, 12, bold=True)

    # Reading note 10pt italics (if present)
    if reading_note and reading_note.strip():
        p_note = doc.add_paragraph()
        _set_para_base(p_note)
        _add_run(p_note, reading_note.strip(), font_name, 10, italic=True)

    # Small gap after header area
    doc.add_paragraph("")

    # Prepare chapter insertion pointers per marker kind
    chapters_by_kind = {"Page": [], "Location": []}
    for ch in chapters:
        if ch.marker_kind in chapters_by_kind:
            chapters_by_kind[ch.marker_kind].append(ch)

    for k in chapters_by_kind:
        chapters_by_kind[k].sort(key=lambda x: x.marker_value)

    next_idx = {"Page": 0, "Location": 0}

    def maybe_insert_chapter(kind: Optional[str], val: Optional[int]):
        if not kind or val is None:
            return
        if kind not in chapters_by_kind:
            return

        i = next_idx[kind]
        lst = chapters_by_kind[kind]
        # Insert all chapter headings whose threshold is <= current marker value
        while i < len(lst) and val >= lst[i].marker_value:
            p_ch = doc.add_paragraph()
            _set_para_base(p_ch)
            _add_run(p_ch, lst[i].title.strip(), font_name, 11, bold=True)
            doc.add_paragraph("")  # small gap after chapter heading
            i += 1
        next_idx[kind] = i

    for e in entries:
        maybe_insert_chapter(e.marker_kind, e.marker_value)

        # Marker line (bold 10pt)
        if e.marker_kind and e.marker_value is not None:
            pm = doc.add_paragraph()
            _set_para_base(pm)
            _add_run(pm, f"{e.marker_kind} {e.marker_value}", font_name, 10, bold=True)

        # Highlight (10pt)
        ph = doc.add_paragraph()
        _set_para_base(ph)
        _add_run(ph, e.highlight, font_name, 10)

        # Note: bullet, NO INDENT, 10pt; only "Note:" bold
        if e.note:
            pbn = doc.add_paragraph()
            _set_para_base(pbn)
            _add_run(pbn, "• ", font_name, 10)
            _add_run(pbn, "Note:", font_name, 10, bold=True)
            _add_run(pbn, f" {e.note}", font_name, 10)

        # Separator line BETWEEN entries only
        ps = doc.add_paragraph()
        _set_para_base(ps)
        _add_run(ps, "-" * 48, font_name, 10)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
